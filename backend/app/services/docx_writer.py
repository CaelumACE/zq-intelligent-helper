"""Sprint 04 公文 docx 导出（国标公文格式）。

与 S02 的 docx_exporter.py（差异报告）相互独立，本模块仅负责将
对话式公文写作生成的 Markdown 正文渲染为符合国标排版的 Word 文档。
"""
import io
import os
import re
from functools import lru_cache
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# 国标页边距（GB/T 9704）：上 3.7cm 下 3.5cm 左 2.8cm 右 2.6cm
PAGE_MARGIN_TOP_CM = 3.7
PAGE_MARGIN_BOTTOM_CM = 3.5
PAGE_MARGIN_LEFT_CM = 2.8
PAGE_MARGIN_RIGHT_CM = 2.6

# 行距 28 磅（固定值）
LINE_SPACING_PT = 28

# 正文字号：三号 = 16 磅；标题字号：二号 = 22 磅
BODY_FONT_SIZE_PT = 16
TITLE_FONT_SIZE_PT = 22

# 字体回退链（按优先级，不硬编码字体文件路径）
FANGSONG_FALLBACK = ("仿宋_GB2312", "仿宋", "FangSong", "STFangsong", "SimSun", "宋体", "serif")
TITLE_FONT_FALLBACK = ("方正小标宋简体", "小标宋体", "小标宋", "SimSun", "宋体", "serif")
HEITI_FALLBACK = ("黑体", "SimHei", "sans-serif")

# 常见系统字体目录，用于探测字体是否安装（仅按文件名匹配，不硬编码具体路径）
_FONT_DIRS = (
    "/usr/share/fonts",
    "/usr/local/share/fonts",
    os.path.expanduser("~/.fonts"),
    os.path.expanduser("~/Library/Fonts"),
    "/Library/Fonts",
    "C:/Windows/Fonts",
    r"C:\Windows\Fonts",
)


@lru_cache(maxsize=1)
def _installed_font_files() -> List[str]:
    """扫描系统字体目录，返回已安装字体文件名列表（小写）。失败时返回空列表。"""
    names: List[str] = []
    for d in _FONT_DIRS:
        try:
            p = Path(d)
            if not p.is_dir():
                continue
            for f in p.rglob("*"):
                if f.is_file() and f.suffix.lower() in (".ttf", ".ttc", ".otf"):
                    names.append(f.name.lower())
        except Exception:
            continue
    return names


def _pick_font(candidates: Iterable[str]) -> str:
    """从候选字体名中挑选系统已安装的第一个；都不可用时返回首个候选名（交由 Word 回退）。"""
    candidates = tuple(candidates)
    if not candidates:
        return "serif"
    installed = _installed_font_files()
    if not installed:
        return candidates[0]
    for cand in candidates:
        key = cand.strip().lower()
        if not key:
            continue
        for fname in installed:
            if key in fname:
                return cand
    return candidates[0]


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False):
    """设置 run 的中英文字体、字号与粗体（中文字体需写 eastAsia 才生效）。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _set_line_spacing(paragraph, pt: float = LINE_SPACING_PT):
    """设置段落固定行距（磅）。"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pt)


def _set_first_line_indent(paragraph, chars: int = 2, size_pt: float = BODY_FONT_SIZE_PT):
    """设置首行缩进 N 字符（按字号磅值换算，同时写 firstLineChars 兼容 Word）。"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = pPr.makeelement(qn("w:ind"), {})
        pPr.append(ind)
    # firstLineChars 单位为百分之一字符
    ind.set(qn("w:firstLineChars"), str(int(chars * 100)))
    ind.set(qn("w:firstLine"), str(int(chars * size_pt * 20)))  # twips
    ind.set(qn("w:leftChars"), "0")
    ind.set(qn("w:left"), "0")


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_styled_runs(paragraph, text: str, font_name: str, size_pt: float, base_bold: bool = False):
    """解析 **加粗** 标记并按片段添加 run。"""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _set_run_font(run, font_name, size_pt, bold=base_bold)
        run = paragraph.add_run(m.group(1))
        _set_run_font(run, font_name, size_pt, bold=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, font_name, size_pt, bold=base_bold)


def _strip_prefix(text: str) -> str:
    """去掉 LLM 写作前置说明语，保留公文正文。"""
    t = (text or "").lstrip()
    # 去掉固定的草稿提示语及其后空白行
    t = re.sub(r"^\s*以下为参考草稿[^\n]*\n*", "", t)
    return t


def _extract_title(lines: List[str], title_hint: str = "") -> str:
    """从 Markdown 行中提取标题：优先 # 标题，其次 title_hint。"""
    for line in lines:
        s = line.strip()
        if s.startswith("#"):
            title = s.lstrip("#").strip()
            if title:
                return title
    return (title_hint or "").strip()


def build_writing_docx(content: str, title_hint: str = "") -> bytes:
    """将 Markdown 格式的公文内容渲染为国标格式 docx，返回字节流。

    :param content: 公文 Markdown 正文（# 标题、普通段落、**加粗** 等）
    :param title_hint: 前端或文件名提示标题，当正文无 # 标题时使用
    """
    body_text = _strip_prefix(content)
    raw_lines = body_text.split("\n")

    title = _extract_title(raw_lines, title_hint)
    fangsong = _pick_font(FANGSONG_FALLBACK)
    title_font = _pick_font(TITLE_FONT_FALLBACK)

    doc = Document()

    # 页边距（国标）
    for section in doc.sections:
        section.top_margin = Cm(PAGE_MARGIN_TOP_CM)
        section.bottom_margin = Cm(PAGE_MARGIN_BOTTOM_CM)
        section.left_margin = Cm(PAGE_MARGIN_LEFT_CM)
        section.right_margin = Cm(PAGE_MARGIN_RIGHT_CM)

    # 设置 Normal 样式默认字体
    try:
        normal = doc.styles["Normal"]
        normal.font.name = fangsong
        normal.font.size = Pt(BODY_FONT_SIZE_PT)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), fangsong)
    except Exception:
        pass

    # 标题：二号小标宋体，居中加粗
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_spacing(p, pt=TITLE_FONT_SIZE_PT * 1.5)
        run = p.add_run(title)
        _set_run_font(run, title_font, TITLE_FONT_SIZE_PT, bold=True)

    # 收集正文段落（跳过已渲染的 # 标题行）
    paragraphs: List[str] = []
    title_consumed = False
    for line in raw_lines:
        s = line.strip()
        if not s:
            paragraphs.append("")
            continue
        if s.startswith("#"):
            if not title_consumed:
                title_consumed = True
                continue
            # 次级标题（## 等）作为正文小标题，黑体
            paragraphs.append("##" + s.lstrip("#").strip())
            continue
        paragraphs.append(s)

    # 落款/日期右对齐：连续出现在文末的、非主送机关的短行
    n = len(paragraphs)
    sign_start = n
    for i in range(n - 1, -1, -1):
        s = paragraphs[i].strip()
        if not s:
            if sign_start - i > 2:
                break
            continue
        # 日期特征（含年月日/〇/数字）或落款单位特征
        looks_like_sign = bool(
            re.search(r"[0-9〇一二三四五六七八九十]{4}\s*年", s)
            or s.endswith(("局", "委", "办", "厅", "部", "署", "院", "会", "府", "处", "科", "公司", "中心", "站"))
        )
        if looks_like_sign:
            sign_start = i
            continue
        break

    for idx, para_text in enumerate(paragraphs):
        s = para_text.strip()
        if not s:
            continue

        p = doc.add_paragraph()
        _set_line_spacing(p)

        if s.startswith("##"):
            # 小标题：黑体加粗，不缩进
            run = p.add_run(s[2:].strip())
            _set_run_font(run, _pick_font(HEITI_FALLBACK), BODY_FONT_SIZE_PT, bold=True)
        elif idx >= sign_start:
            # 落款 / 日期：右对齐，不缩进
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_styled_runs(p, s, fangsong, BODY_FONT_SIZE_PT)
        elif _is_recipient(s):
            # 主送机关：三号仿宋，顶格
            _add_styled_runs(p, s, fangsong, BODY_FONT_SIZE_PT)
        else:
            # 正文：三号仿宋，首行缩进 2 字符
            _set_first_line_indent(p, chars=2, size_pt=BODY_FONT_SIZE_PT)
            _add_styled_runs(p, s, fangsong, BODY_FONT_SIZE_PT)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


_RECIPIENT_SUFFIX = ("：", ":")
_RECIPIENT_KEYWORDS = ("单位", "机关", "部门", "公司", "各", "同志")


def _is_recipient(line: str) -> bool:
    """判断是否为主送机关行：以冒号结尾且较短，或含机关单位特征词。"""
    if not line or len(line) > 40:
        return False
    if line.endswith(_RECIPIENT_SUFFIX) and "\n" not in line:
        # 排除标题层级
        if line.startswith("#"):
            return False
        return any(k in line for k in _RECIPIENT_KEYWORDS) or line.endswith("：")
    return False
