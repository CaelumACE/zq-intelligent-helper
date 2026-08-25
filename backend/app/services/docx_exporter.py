"""Sprint 02 差异报告 docx 导出。"""
import io

from docx import Document
from docx.shared import RGBColor


def build_diff_docx(doc_a_title: str, doc_b_title: str, result: dict) -> bytes:
    doc = Document()
    doc.add_heading("政策差异比对报告", 0)

    info = doc.add_paragraph()
    info.add_run("旧版文档：").bold = True
    info.add_run(doc_a_title)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("新版文档：").bold = True
    p.add_run(doc_b_title)

    summary = result.get("summary") or {}
    doc.add_heading("变更摘要", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["新增", "删除", "修改", "变更总数"]
    values = [summary.get("added", 0), summary.get("removed", 0), summary.get("modified", 0), summary.get("total_changes", 0)]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for i, v in enumerate(values):
        table.rows[1].cells[i].text = str(v)
    if summary.get("brief"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("变更方向：").bold = True
        p.add_run(str(summary.get("brief")))

    doc.add_heading("逐条对比", level=1)
    diffs = result.get("diffs") or []
    if not diffs:
        doc.add_paragraph("未发现差异。")
    for i, d in enumerate(diffs, 1):
        label = {"added": "新增", "removed": "删除", "modified": "修改"}.get(d.get("type"), "变更")
        doc.add_heading(f"{i}. [{label}] {d.get('clause') or ''}", level=2)
        if d.get("old_text"):
            row = doc.add_paragraph()
            row.add_run("旧版：").bold = True
            run = row.add_run(d.get("old_text", ""))
            if d.get("type") == "removed":
                run.font.strike = True
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        if d.get("new_text"):
            row = doc.add_paragraph()
            row.add_run("新版：").bold = True
            run = row.add_run(d.get("new_text", ""))
            if d.get("type") == "added":
                run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)
        if d.get("change_note"):
            row = doc.add_paragraph()
            row.add_run("说明：").bold = True
            row.add_run(d.get("change_note", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
